# generar_documento_word.py
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from pathlib import Path

BASE_DIR = Path("D:/Proyectos/P089 - Catastro Compras Cambio Climatico")
FIG_DIR = BASE_DIR / "figuras_catastro"

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def create_word_report(output_path):
    doc = docx.Document()
    
    # Configurar márgenes estándar
    sections = doc.sections
    for s in sections:
        s.top_margin = Inches(1)
        s.bottom_margin = Inches(1)
        s.left_margin = Inches(1)
        s.right_margin = Inches(1)
        
    COLOR_PRIMARY = RGBColor(0x1F, 0x49, 0x7D)    # Azul institucional
    COLOR_SECONDARY = RGBColor(0x36, 0x60, 0x92)  # Azul medio
    COLOR_DARK = RGBColor(0x26, 0x26, 0x26)       # Gris oscuro texto
    COLOR_ALERT = RGBColor(0x9C, 0x00, 0x06)      # Rojo advertencia
    
    # -------------------------------------------------------------
    # PORTADA / ENCABEZADO PRINCIPAL
    # -------------------------------------------------------------
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("Catastro Histórico de Compras Públicas en Cambio Climático de Chile (2007–2026)")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(19)
    run_title.font.bold = True
    run_title.font.color.rgb = COLOR_PRIMARY
    
    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(8)
    run_sub = sub_p.add_run("Informe Metodológico, Radiografía de Mecanismos de Compra, Gobernanza Multinivel y Codebook\nElaborado por: Eduardo Vega Toledo | A encargo de: Valentina Cariaga Cerda")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(10.5)
    run_sub.font.italic = True
    run_sub.font.color.rgb = COLOR_SECONDARY
    
    # Tabla de metadata de autoría
    t_meta = doc.add_table(rows=5, cols=2)
    t_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Elaborado por:", "Eduardo Vega Toledo (Administrador Público | M.Sc. | Data Engineer)"),
        ("A encargo de:", "Valentina Cariaga Cerda"),
        ("Cobertura y Fuentes:", "Mercado Público (ChileCompra 2007 a julio 2026 — 470 bases masivas)"),
        ("Universo Analizado:", "9.086 procesos únicos deduplicados ($359.604 millones de pesos transados)"),
        ("Contacto y Licencia:", "evega.ap@gmail.com | Creative Commons Attribution 4.0 (CC BY 4.0)")
    ]
    for row_idx, (k, v) in enumerate(meta_data):
        row = t_meta.rows[row_idx]
        cell_k, cell_v = row.cells[0], row.cells[1]
        
        cell_k.width = Inches(2.2)
        cell_v.width = Inches(4.3)
        set_cell_background(cell_k, "F2F2F2")
        set_cell_background(cell_v, "FFFFFF")
        set_cell_margins(cell_k, 70, 70, 90, 90)
        set_cell_margins(cell_v, 70, 70, 90, 90)
        
        pk = cell_k.paragraphs[0]
        pk.paragraph_format.space_after = Pt(0)
        rk = pk.add_run(k)
        rk.font.name = "Arial"
        rk.font.size = Pt(9)
        rk.font.bold = True
        rk.font.color.rgb = COLOR_PRIMARY
        
        pv = cell_v.paragraphs[0]
        pv.paragraph_format.space_after = Pt(0)
        rv = pv.add_run(v)
        rv.font.name = "Arial"
        rv.font.size = Pt(9)
        rv.font.color.rgb = COLOR_DARK
        
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    
    def add_h1(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(4)
        r = h.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(12.5)
        r.font.bold = True
        r.font.color.rgb = COLOR_PRIMARY
        return h

    def add_h2(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(10)
        h.paragraph_format.space_after = Pt(3)
        r = h.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(10.5)
        r.font.bold = True
        r.font.color.rgb = COLOR_SECONDARY
        return h

    def add_body(text, bold_prefix="", italic=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            rb = p.add_run(bold_prefix)
            rb.font.name = "Arial"
            rb.font.size = Pt(9.5)
            rb.font.bold = True
            rb.font.color.rgb = COLOR_DARK
        r = p.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(9.5)
        r.font.italic = italic
        r.font.color.rgb = COLOR_DARK
        return p

    def add_figure(fig_path, caption_text, width_inches=6.0):
        if Path(fig_path).exists():
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_before = Pt(6)
            p_img.paragraph_format.space_after = Pt(2)
            run = p_img.add_run()
            run.add_picture(str(fig_path), width=Inches(width_inches))
            
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_after = Pt(8)
            r_cap = p_cap.add_run(caption_text)
            r_cap.font.name = "Arial"
            r_cap.font.size = Pt(8.5)
            r_cap.font.italic = True
            r_cap.font.color.rgb = COLOR_DARK

    # -------------------------------------------------------------
    # 1. RESUMEN EJECUTIVO Y DINÁMICA TEMPORAL
    # -------------------------------------------------------------
    add_h1("1. Resumen Ejecutivo y Dinámica Temporal (2007–2026)")
    add_body("El presente catastro consolida la totalidad de compras públicas y licitaciones del Estado de Chile orientadas a cambio climático, adaptación y mitigación a lo largo de 20 años continuos (2007 a julio de 2026). El corpus reúne 9.086 procesos de compra únicos deduplicados (2.287 licitaciones públicas y 6.799 órdenes de compra), transando un monto acumulado de $359.604 millones de pesos chilenos (~USD 380 millones).")
    add_body("El análisis longitudinal evidencia un quiebre estructural a partir de la promulgación de la Ley Marco de Cambio Climático (Ley 21.455 en junio de 2022). De un promedio histórico de 350 a 450 compras anuales durante el período 2007-2021, las contrataciones escalaron a 777 en 2023 (+55%) y alcanzaron su máximo histórico en 2024 con 1.027 procesos (+105% respecto a 2022), transando más de $111.506 millones de pesos en un solo año.")

    add_figure(FIG_DIR / "fig1_evolucion_temporal_ley21455.png",
               "Figura 1: Evolución de las compras climáticas y quiebre estructural tras la Ley Marco 21.455 (2007–2026).")

    # -------------------------------------------------------------
    # 2. MECANISMOS DE COMPRA: LICITACIONES VS OC Y TRATO DIRECTO
    # -------------------------------------------------------------
    add_h1("2. Mecanismos de Contratación: Licitación, Convenio Marco y Trato Directo")
    add_body("En la gestión pública chilena, una Licitación y una Orden de Compra NO son equivalentes y no deben confundirse:")
    add_body("Es el concurso público administrativo abierto que define bases técnicas, competencia de oferentes y selección de proveedor. Registra los grandes proyectos y llamados estratégicos.", "• Licitación Pública (2.287 procesos): ")
    add_body("Es el documento transaccional de despacho y pago vinculante. De las 6.799 OCs, más de la mitad NO derivan de licitaciones locales, sino de mecanismos directos:", "• Órdenes de Compra (6.799 transacciones): ")

    add_figure(FIG_DIR / "fig2_mecanismos_contratacion.png",
               "Figura 2: Distribución por Mecanismo de Contratación (Licitaciones, Convenio Marco, Trato Directo y Compra Ágil).")

    t_mec = doc.add_table(rows=1, cols=4)
    t_mec.alignment = WD_TABLE_ALIGNMENT.CENTER
    mec_headers = ["Mecanismo Legal de Contratación", "N° Procesos", "Monto (Millones CLP)", "% Participación"]
    for i, h in enumerate(mec_headers):
        cell = t_mec.rows[0].cells[i]
        set_cell_background(cell, "1F497D")
        set_cell_margins(cell, 50, 50, 60, 60)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(h)
        r.font.name = "Arial"
        r.font.size = Pt(8.5)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
    mec_rows = [
        ("Licitación Pública (Concursos Abiertos)", "2.250", "$242.277M", "24,8%"),
        ("Trato Directo (Mecanismo Excepcional)", "3.417", "$24.476M", "37,6%"),
        ("Convenio Marco (Catálogo Electrónico)", "2.279", "$3.671M", "25,1%"),
        ("Compra Ágil (Menor a 30 UTM)", "999", "$745M", "11,0%"),
        ("Licitación Privada (Concurso Cerrado)", "37", "$139M", "0,4%"),
        ("Otras Órdenes de Compra", "104", "$894M", "1,1%"),
        ("TOTAL", "9.086", "$359.604M", "100.0%")
    ]
    for r_data in mec_rows:
        row = t_mec.add_row()
        is_total = "TOTAL" in r_data[0]
        bg = "EBF1F5" if is_total else "FFFFFF"
        for i, val in enumerate(r_data):
            cell = row.cells[i]
            set_cell_background(cell, bg)
            set_cell_margins(cell, 40, 40, 60, 60)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(val)
            r.font.name = "Arial"
            r.font.size = Pt(8)
            r.font.bold = is_total
            r.font.color.rgb = COLOR_PRIMARY if is_total else COLOR_DARK

    add_h2("La Relevancia del Trato Directo en Cambio Climático")
    add_body("El 37,6% de todas las compras del catastro (3.417 órdenes de compra por $24.476 millones) fueron suscritas vía Trato Directo (mecanismo excepcional). En ciencia política y gestión pública, esto refleja situaciones de emergencia climática (aluviones, incendios forestales, sequía extrema decretada en comunas), contrataciones de urgencia de camiones aljibe y consultorías con proveedores técnicos exclusivos.")

    # -------------------------------------------------------------
    # 3. GOBERNANZA MULTINIVEL
    # -------------------------------------------------------------
    add_h1("3. Radiografía de Gobernanza Multinivel (Municipalidades, GOREs y Central)")
    add_body("El catastro incorpora una clasificación institucional que permite analizar la descentralización de las compras y la distribución de capacidades presupuestarias entre niveles del Estado:")

    add_figure(FIG_DIR / "fig3_gobernanza_multinivel.png",
               "Figura 3: Distribución de compras climáticas y montos transados por Nivel Institucional (2007–2026).")

    add_h2("A. El Nivel Municipal (2.605 compras y 54,8% del gasto licitado)")
    add_body("Las municipalidades representan más de la mitad del presupuesto licitado debido a grandes obras comunales de infraestructura verde, recambio de luminarias LED de alta eficiencia energética, arbolado y consultorías de Planes de Acción Comunal de Cambio Climático (PACCC). Municipios destacados: Concepción (127 compras), Copiapó (91), Lampa (66), El Bosque (61) y Temuco (60).")

    add_h2("B. Los Gobiernos Regionales (142 compras en escala mesoterritorial)")
    add_body("Los GOREs concentran contrataciones estratégicas orientadas a los Planes de Acción Regional de Cambio Climático (PARCC), balances y modelos de cuencas hídricas y estudios de vulnerabilidad regional. Destacan GORE O'Higgins (18 compras / $785M), GORE Metropolitano (17 compras / $838M), GORE Atacama (13 compras / $238M) y GORE Coquimbo (15 compras).")

    add_figure(FIG_DIR / "fig5_top_organismos_compradores.png",
               "Figura 4: Top 12 instituciones compradoras en cambio climático del Estado de Chile.")

    # -------------------------------------------------------------
    # 4. LIMITACIONES, RIESGOS Y CONTROL DE CALIDAD
    # -------------------------------------------------------------
    add_h1("4. Límites Epistemológicos, Restricciones y Control de Calidad")
    add_body("Para asegurar la máxima rigurosidad y confiabilidad en el uso de estos datos en investigaciones, evaluaciones o políticas públicas, se deben considerar explícitamente los siguientes alcances y restricciones:")
    
    add_body("La emisión de una orden de compra o la adjudicación de una licitación prueba fehacientemente la asignación presupuestaria y la priorización administrativa del organismo, pero no equivale automáticamente a la evaluación de impacto ecológico posterior ni a la ejecución física final de la obra.", "1. Intención de Compra vs. Impacto Real: ")
    
    add_body("Mercado Público no cuenta con un etiquetado nativo para cambio climático. La clasificación depende de la redacción técnica de los funcionarios de compras. Se mitigan omisiones mediante escaneo exhaustivo de 8 campos textuales (nombre, descripción, especificaciones técnicas y rubros).", "2. Sesgo de Rotulado Administrativo: ")
    
    add_body("Para evitar la sobrestimación del gasto, se debe distinguir estrictamente entre Licitaciones (LP/LE, que reflejan montos marco o estimados) y Órdenes de Compra (OC, que reflejan montos netos transados línea a línea). El dataset incluye la variable 'tipo_registro' y 'mecanismo_compra' para filtrar y evitar doble contabilización.", "3. Riesgo de Doble Contabilización de Montos: ")
    
    add_body("Se aplicó un filtro negativo estricto para eliminar compras espurias de 'clima laboral', 'ambiente laboral' o 'aire acondicionado', así como siglas ambiguas como SBN (subvención municipal no relacionada).", "4. Control de Falsos Positivos: ")

    add_body("El dataset final superó 16 de 16 pruebas formales de auditoría informática: cero duplicados en clave compuesta, 100% de campos obligatorios poblados, cuadratura matemática exacta entre pestañas y trazabilidad a través del enlace web público a cada proceso.", "5. Certificación de Calidad y Consistencia: ")

    # -------------------------------------------------------------
    # 5. TAXONOMÍA TEMÁTICA Y ECOSISTEMA DE PROVEEDORES
    # -------------------------------------------------------------
    add_h1("5. Taxonomía Temática y Ecosistema de Proveedores")
    add_body("El catastro clasifica las 9.086 compras en 4 subcategorías conceptuales:")

    add_figure(FIG_DIR / "fig4_subcategorias_tematicas.png",
               "Figura 5: Distribución de compras por Eje Temático y Modalidad de Registro.")

    add_body("Empresas de ingeniería ambiental como Deuman (69 contratos adjudicados por $1.765 millones) lideran los estudios técnicos de inventarios GEI, huella de carbono y hojas de ruta de descarbonización e hidrógeno verde.", "• Consultoría Especializada: ")
    add_body("La Universidad de Chile (88 adjudicaciones por $1.380M) y la Pontificia Universidad Católica de Chile (55 adjudicaciones por $644M) actúan como asesores científicos principales del Ministerio del Medio Ambiente y Energía.", "• Academia y Validación Científica: ")
    add_body("La Asociación Chilena de Municipalidades (AChM, 105 contrataciones) cumple un rol articulador entregando asistencia técnica y capacitación para ordenanzas y planes comunales.", "• Asociativismo y Soporte Local: ")

    # -------------------------------------------------------------
    # 6. CODEBOOK OPERACIONAL
    # -------------------------------------------------------------
    add_h1("6. Codebook / Diccionario de Variables")
    add_body("La base de datos entregada en Excel y CSV contiene 23 variables operacionales:")

    t_code = doc.add_table(rows=1, cols=3)
    t_code.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_headers = ["Variable", "Tipo", "Descripción y Valores de Ejemplo"]
    for i, h in enumerate(c_headers):
        cell = t_code.rows[0].cells[i]
        set_cell_background(cell, "1F497D")
        set_cell_margins(cell, 50, 50, 60, 60)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.name = "Arial"
        r.font.size = Pt(8.5)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
    code_rows = [
        ("codigo_proceso", "Texto", "Identificador único de la licitación (CodigoExterno) o de la orden de compra."),
        ("tipo_registro", "Categórica", "Modalidad de registro: 'licitacion' o 'orden_compra'."),
        ("mecanismo_compra", "Categórica", "Licitación Pública, Convenio Marco, Trato Directo (Excepcional), Compra Ágil (<30 UTM), Licitación Privada."),
        ("link", "URL", "Enlace directo a la ficha pública en Mercado Público para verificación."),
        ("nombre", "Texto", "Nombre formal del proceso de licitación o de la orden de compra."),
        ("descripcion", "Texto", "Glosa descriptiva detallada del requerimiento técnico."),
        ("organismo_comprador", "Texto", "Nombre de la institución pública compradora."),
        ("nivel_institucional", "Categórica", "Municipalidades, Gobiernos Regionales (GORE), Gobierno Central, Universidades, Salud, Defensa."),
        ("rut_comprador", "Texto", "RUT institucional del organismo comprador."),
        ("region_comprador", "Categórica", "Región político-administrativa del comprador (16 regiones)."),
        ("fecha", "Fecha", "Fecha de publicación o emisión (2007-01 a 2026-07)."),
        ("monto_pesos", "Numérico", "Monto en pesos chilenos estimado/adjudicado o transado."),
        ("moneda", "Categórica", "Moneda original de la contratación (CLP, UF, USD, UTM)."),
        ("proveedor", "Texto", "Razón social del proveedor adjudicado."),
        ("rut_proveedor", "Texto", "RUT del proveedor adjudicado."),
        ("subcategoria", "Categórica", "Núcleo_Exacto, Instrumentos_Ley21455, Gases_Descarbonizacion, Transicion_SBN."),
        ("termino_coincidente", "Texto", "Palabra o frase exacta que gatilló la clasificación positiva."),
        ("texto_fragmento", "Texto", "Fragmento de texto de 80 caracteres con el contexto del término.")
    ]
    for r_data in code_rows:
        row = t_code.add_row()
        for i, val in enumerate(r_data):
            cell = row.cells[i]
            set_cell_background(cell, "FFFFFF")
            set_cell_margins(cell, 35, 35, 50, 50)
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.name = "Arial"
            r.font.size = Pt(8)
            if i == 0:
                r.font.bold = True
                r.font.color.rgb = COLOR_PRIMARY
            else:
                r.font.color.rgb = COLOR_DARK

    # -------------------------------------------------------------
    # 7. PROTOCOLO DE REPRODUCIBILIDAD Y SCRIPTS EN PYTHON
    # -------------------------------------------------------------
    add_h1("7. Protocolo de Reproducibilidad y Scripts en Python")
    add_body("El dataset y las visualizaciones fueron construidos bajo una arquitectura determinista de código abierto. Se incluyen dos scripts en Python en la carpeta de entrega:")
    add_body("Script autónomo y comentado paso a paso para que cualquier persona ejecute la extracción sobre las bases masivas de ChileCompra y replique exactamente los 9.086 registros.", "1. replicar_catastro_chilecompra.py: ")
    add_body("Script ejecutable para regenerar automáticamente todas las figuras y gráficas estadísticas (300 DPI) a partir del archivo CSV.", "2. generar_graficas_catastro.py: ")

    # -------------------------------------------------------------
    # 8. PREGUNTAS DE INVESTIGACIÓN Y PROPUESTAS METODOLÓGICAS
    # -------------------------------------------------------------
    add_h1("8. Agendas de Investigación y Diseños Metodológicos Sugeridos")
    add_body("A partir de este conjunto de datos empíricos, se formulan 4 agendas de investigación científica y evaluación de políticas públicas, detallando los modelos cuantitativos y estrategias de cruce recomendadas:")

    add_h2("Agenda 1: Capacidad Estatal y Brechas Territoriales en Gobiernos Locales")
    add_body("¿En qué medida los ingresos propios comunales, la dependencia del FCM y la dotación técnica condicionan la probabilidad y celeridad de licitar instrumentos climáticos (PACCC)?", "• Pregunta: ")
    add_body("Modelos de elección discreta (Logit/Probit) para adopción binaria; regresiones Tobit/Poisson para montos per cápita licitados; y modelos de supervivencia (Cox Proportional Hazards) para medir el tiempo transcurrido hasta la primera licitación post-Ley 21.455. Cruce de datos sugerido con SINIM (SUBDERE), CASEN (pobreza multidimensional) y Censo.", "• Diseño Metodológico: ")

    add_h2("Agenda 2: Evaluación de Impacto Normativo (Efecto Causal de la Ley 21.455)")
    add_body("¿Cuál es el impacto causal de la Ley Marco de Cambio Climático sobre la aceleración del gasto y la descentralización de compras ambientales?", "• Pregunta: ")
    add_body("Diseño de Series Temporales Interrumpidas (ITSA) mensual comparando trayectorias pre-2022 vs post-2022, complementado con Diferencias en Diferencias (DiD) entre organismos directamente obligados vs organismos no regulados.", "• Diseño Metodológico: ")

    add_h2("Agenda 3: Redes de Política Pública y Concentración de Proveedores (Policy Networks)")
    add_body("¿Cómo se estructura la red de gobernanza público-privada y cuán concentrado está el mercado de consultoría frente a la academia y gremios?", "• Pregunta: ")
    add_body("Análisis de Redes Sociales Bipartitas (Two-Mode Network Analysis) entre compradores y proveedores con métricas de intermediación (Betweenness), autovalor (Eigenvector) y algoritmos de comunidades (Louvain). Cálculo de Índice de Herfindahl-Hirschman (HHI) por submercados.", "• Diseño Metodológico: ")

    add_h2("Agenda 4: Análisis Textual y Calidad de la Demanda Estatal (NLP & Topic Modeling)")
    add_body("¿Qué prioriza materialmente el Estado: diagnósticos institucionales (soft governance) u obras de infraestructura resiliente (hard adaptation)?", "• Pregunta: ")
    add_body("Modelado Estructural de Tópicos (Structural Topic Model - STM / BERTopic) sobre glosas textuales con covariables institucionales, y clasificación supervisada de intervenciones tangibles vs intangibles.", "• Diseño Metodológico: ")

    doc.save(str(output_path))
    print(f"Documento Word guardado exitosamente: {output_path}")

if __name__ == "__main__":
    out1 = Path(r"D:\Proyectos\P089 - Catastro Compras Cambio Climatico\Catastro_Compras_Cambio_Climatico_Informe_Metodologico_y_Estadistico.docx")
    out2 = Path(r"C:\Users\evega\OneDrive\Documents\Obsidian\MyWorld\2 - Project\P089 - Catastro Compras Cambio Climatico\Catastro Compras Cambio Climatico 2007-2026 V.0 25082026\Catastro_Compras_Cambio_Climatico_Informe_Metodologico_y_Estadistico.docx")
    create_word_report(out1)
    create_word_report(out2)
